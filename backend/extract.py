"""
Universal file reader: any file -> records (list of {column: value}).

Structured files (CSV/TSV/JSON/Excel) are parsed locally -- exact and free.
Unstructured files (PDF/Word/text/markdown) are turned into structured rows by
the LLM: we pull the text out locally, then ask Claude to extract the most
relevant table as JSON. That is the "ingest any file" path.

Returns: {records, headers, kind, note, cost_usd}
"""
import csv
import io
import json
import os

import llm

MAX_DOC_CHARS = 16000   # keep LLM extraction prompt bounded


def _from_rows(header, rows):
    header = [str(h).strip() if h is not None else "col_%d" % i for i, h in enumerate(header)]
    records = []
    for r in rows:
        if all(c is None or str(c).strip() == "" for c in r):
            continue
        records.append({header[i]: r[i] if i < len(r) else None for i in range(len(header))})
    return header, records


def _csv(path, delim=","):
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        rows = list(csv.reader(f, delimiter=delim))
    if not rows:
        return [], []
    return _from_rows(rows[0], rows[1:])


def _xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    if not rows:
        return [], []
    return _from_rows(rows[0], rows[1:])


def _json(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        data = json.load(f)
    if isinstance(data, list) and data and isinstance(data[0], dict):
        recs = data
    elif isinstance(data, dict):
        # find the first list-of-dicts inside
        recs = None
        for v in data.values():
            if isinstance(v, list) and v and isinstance(v[0], dict):
                recs = v
                break
        if recs is None:
            recs = [data]
    else:
        recs = [{"value": data}]
    headers = list(recs[0].keys()) if recs else []
    return headers, recs


def _pdf_text(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join((pg.extract_text() or "") for pg in reader.pages)


def _docx_text(path):
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _txt(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _llm_extract(text, filename):
    """Ask the LLM to pull the most relevant table out of unstructured text."""
    if not llm.available():
        raise RuntimeError("This file type needs LLM extraction, but no LLM backend "
                           "is configured (set MAGICA_API_KEY in .env).")
    text = text[:MAX_DOC_CHARS]
    sys = "You extract structured tabular data from business documents for a data pipeline. Output strict JSON only, no prose."
    prompt = (
        "From the document below, extract the most relevant structured data as records.\n"
        "Return JSON exactly as: {\"headers\": [\"col1\", ...], \"records\": [{\"col1\": value, ...}, ...]}\n"
        "- Produce a WIDE table: ONE row per entity (per company / market / product), with a\n"
        "  separate named COLUMN for each distinct measure. Do NOT melt into a long\n"
        "  'Metric'/'Value' shape -- e.g. for a market note, return a single row with columns\n"
        "  like tam, sam, cagr_value, cagr_volume, hhi, avg_firm_age, not rows of metric/value pairs.\n"
        "- Use short, clear snake_case column names taken from the document.\n"
        "- Keep values as written (don't normalise units or currency).\n"
        "- If there are several tables, pick the largest / most business-relevant one.\n\n"
        "DOCUMENT (" + filename + "):\n" + text
    )
    res = llm.complete(prompt, system=sys, model=llm.EXTRACT_MODEL, max_tokens=4000)
    parsed = llm.extract_json(res["text"])
    records = parsed.get("records", []) if isinstance(parsed, dict) else parsed
    headers = parsed.get("headers") if isinstance(parsed, dict) else None
    if not headers and records:
        headers = list(records[0].keys())
    return headers or [], records or [], res.get("cost_usd", 0)


def extract(path):
    ext = os.path.splitext(path)[1].lower()
    name = os.path.basename(path)
    cost = 0.0

    if ext in (".csv",):
        headers, records = _csv(path, ",")
        kind, note = "table", "parsed locally as CSV"
    elif ext in (".tsv",):
        headers, records = _csv(path, "\t")
        kind, note = "table", "parsed locally as TSV"
    elif ext in (".xlsx", ".xlsm"):
        headers, records = _xlsx(path)
        kind, note = "table", "parsed locally from Excel"
    elif ext in (".json",):
        headers, records = _json(path)
        kind, note = "table", "parsed locally as JSON"
    elif ext in (".pdf", ".docx", ".txt", ".md", ".markdown"):
        if ext == ".pdf":
            text = _pdf_text(path)
        elif ext == ".docx":
            text = _docx_text(path)
        else:
            text = _txt(path)
        if not text.strip():
            raise RuntimeError("Could not read text from %s (a scanned image PDF would need OCR/vision)." % name)
        headers, records, cost = _llm_extract(text, name)
        kind, note = "doc", "extracted by Claude from an unstructured %s" % ext.lstrip(".")
    else:
        raise RuntimeError("Unsupported file type: %s" % ext)

    return {"records": records, "headers": headers, "kind": kind, "note": note,
            "cost_usd": cost, "filename": name}
